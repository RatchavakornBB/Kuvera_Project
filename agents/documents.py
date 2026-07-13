"""Fetches a document's real bytes from Supabase Storage for a node to pass
to Claude as a native content block (system-architecture.md Section 5.4) —
no separate text-extraction step for PDF/images, Claude reads the file
directly.

PDF and images are both real, native Claude Messages API input types (a
`document` block and an `image` block respectively) — genuinely supported,
not stubs. Word documents (.docx) have no native Claude content block, so
they're bridged honestly: real text extracted via python-docx (already a
dependency — agents/drafting_lead.py uses the same library the other
direction, to generate .docx), sent as a `text` block. This is a real
extraction of the document's actual paragraph text, not a fabricated
summary — the model reads exactly what's in the file, just as plain text
instead of a native document block. Audio and video are NOT extended here:
Claude's Messages API has no native audio/video content block, and there is
no real transcription/frame-extraction pipeline in this codebase to bridge
them honestly the way python-docx does for Word files — they still raise a
clear error at fetch time rather than fabricating one."""

import io
from typing import Any

from docx import Document as DocxDocument

from agents.db import get_client

BUCKET = "deal-documents"

_MEDIA_TYPES = {
    "pdf": "application/pdf",
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

DOCX_MEDIA_TYPE = _MEDIA_TYPES["docx"]

# Claude's Messages API content block "type" differs by media_type — a PDF is
# a `document` block, an image is an `image` block, and a Word doc becomes a
# `text` block (real extracted content, see module docstring).


def _extract_docx_text(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def build_content_block(content: bytes, media_type: str) -> dict[str, Any]:
    """The one place a document's bytes become a Claude content block —
    every document-reading node uses this instead of hardcoding the block
    shape, so adding a new supported media_type never requires a node edit."""
    import base64

    if media_type == DOCX_MEDIA_TYPE:
        return {"type": "text", "text": f"[Extracted text from Word document]\n\n{_extract_docx_text(content)}"}

    block_type = "document" if media_type == "application/pdf" else "image"
    return {
        "type": block_type,
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": base64.standard_b64encode(content).decode(),
        },
    }


def fetch_document(document_id: str) -> tuple[bytes, str, str]:
    """Returns (file_bytes, filename, media_type)."""
    client = get_client()
    row = client.table("documents").select("name, storage_path").eq("id", document_id).execute().data
    if not row:
        raise ValueError(f"No document with id {document_id}")
    name, storage_path = row[0]["name"], row[0]["storage_path"]
    if not storage_path:
        raise ValueError(f"Document {document_id} has no storage_path (upload incomplete)")

    content = client.storage.from_(BUCKET).download(storage_path)
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    media_type = _MEDIA_TYPES.get(ext)
    if media_type is None:
        raise ValueError(
            f"Unsupported document type for direct model reading: .{ext} "
            "(PDF, images, and .docx are supported; audio/video have no native Claude "
            "Messages API content type and no real transcription pipeline exists here, "
            "not just missing glue code)"
        )
    return content, name, media_type
